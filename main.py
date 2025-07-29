import streamlit as st
import requests
import json as js
import numpy as np
import html2text
import google.generativeai as genai
import docx

from markdown import markdown
from xhtml2pdf import pisa
import io
from PyPDF2 import PdfReader

listas_nro_reunioes = np.arange(270, (266 -5 -1), -1)

#---------------------------------------------------------------------------------------------------

# URL do PDF
url_fed = 'https://www.federalreserve.gov/monetarypolicy/files/monetary20250618a1.pdf'\
          #"https://www.federalreserve.gov/monetarypolicy/files/monetary20250129a1.pdf"

# Baixar o PDF
response_fed = requests.get(url_fed)
with open("documento.pdf", "wb") as file:
    file.write(response_fed.content)

# Ler o texto do PDF
reader = PdfReader("documento.pdf")
texto_fed = ""
for page in reader.pages:
    texto_fed += page.extract_text()

#---------------------------------------------------------------------------------------------------

model_ai = genai.GenerativeModel('gemini-1.5-flash')

st.set_page_config("Análise atas Copom",layout="wide")

def markdown_2pdf(markdown_text):
    html_text = markdown(markdown_text)
    pdf_buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(io.StringIO(html_text), dest = pdf_buffer)
    if pisa_status.err:
        return None
    pdf_buffer.seek(0)
    return  pdf_buffer.getvalue()

with st.sidebar:
    nro_ata_selecionada = st.selectbox('Número da ata',listas_nro_reunioes)
    valor_api = st.text_input('API-Key',type = 'password')

    genai.configure(api_key=valor_api)

    # Obtenção das atas do COPOM
    atas = requests.get(
        f'https://www.bcb.gov.br/api/servico/sitebcb/copom/atas_detalhes?nro_reuniao={nro_ata_selecionada}')
    atas_json = js.loads(atas.text)

    texto_ata = js.dumps(atas_json['conteudo'][0]['textoAta'], indent=4)

    # Foramtação do texto da ata
    converter = html2text.HTML2Text()
    converter.ignore_links = True

    atas_formatadas = bytes(texto_ata, "utf-8").decode("unicode_escape")

    atas_formatadas_markdown = converter.handle(atas_formatadas)
    texto_copom = '''
   Copom eleva a taxa Selic para 15,00% a.a. - 271ª reunião - junho 2025
Data de publicação: 18/06/2025

​O ambiente externo mantém-se adverso e particularmente incerto em função da conjuntura e da política econômica nos Estados Unidos, principalmente acerca de suas políticas comercial e fiscal e de seus respectivos efeitos. Além disso, o comportamento e a volatilidade de diferentes classes de ativos também têm sido afetados, com reflexos nas condições financeiras globais. Tal cenário segue exigindo cautela por parte de países emergentes em ambiente de acirramento da tensão geopolítica.

Em relação ao cenário doméstico, o conjunto dos indicadores de atividade econômica e do mercado de trabalho ainda tem apresentado algum dinamismo, mas observa-se certa moderação no crescimento. Nas divulgações mais recentes, a inflação cheia e as medidas subjacentes mantiveram-se acima da meta para a inflação.

As expectativas de inflação para 2025 e 2026 apuradas pela pesquisa Focus permanecem em valores acima da meta, situando-se em 5,2% e 4,5%, respectivamente. A projeção de inflação do Copom para o ano de 2026, atual horizonte relevante de política monetária, situa-se em 3,6% no cenário de referência (Tabela 1).

Os riscos para a inflação, tanto de alta quanto de baixa, seguem mais elevados do que o usual. Entre os riscos de alta para o cenário inflacionário e as expectativas de inflação, destacam-se (i) uma desancoragem das expectativas de inflação por período mais prolongado; (ii) uma maior resiliência na inflação de serviços do que a projetada em função de um hiato do produto mais positivo; e (iii) uma conjunção de políticas econômicas externa e interna que tenham impacto inflacionário maior que o esperado, por exemplo, por meio de uma taxa de câmbio persistentemente mais depreciada. Entre os riscos de baixa, ressaltam-se (i) uma eventual desaceleração da atividade econômica doméstica mais acentuada do que a projetada, tendo impactos sobre o cenário de inflação; (ii) uma desaceleração global mais pronunciada decorrente do choque de comércio e de um cenário de maior incerteza; e (iii) uma redução nos preços das commodities com efeitos desinflacionários.

O Comitê segue acompanhando com atenção como os desenvolvimentos da política fiscal impactam a política monetária e os ativos financeiros. O cenário segue sendo marcado por expectativas desancoradas, projeções de inflação elevadas, resiliência na atividade econômica e pressões no mercado de trabalho. Para assegurar a convergência da inflação à meta em ambiente de expectativas desancoradas, exige-se uma política monetária em patamar significativamente contracionista por período bastante prolongado.

O Copom decidiu elevar a taxa básica de juros em 0,25 ponto percentual, para 15,00% a.a., e entende que essa decisão é compatível com a estratégia de convergência da inflação para o redor da meta ao longo do horizonte relevante. Sem prejuízo de seu objetivo fundamental de assegurar a estabilidade de preços, essa decisão também implica suavização das flutuações do nível de atividade econômica e fomento do pleno emprego.

Em se confirmando o cenário esperado, o Comitê antecipa uma interrupção no ciclo de alta de juros para examinar os impactos acumulados do ajuste já realizado, ainda por serem observados, e então avaliar se o nível corrente da taxa de juros, considerando a sua manutenção por período bastante prolongado, é suficiente para assegurar a convergência da inflação à meta. O Comitê enfatiza que seguirá vigilante, que os passos futuros da política monetária poderão ser ajustados e que não hesitará em prosseguir no ciclo de ajuste caso julgue apropriado.

Votaram por essa decisão os seguintes membros do Comitê: Gabriel Muricca Galípolo (presidente), Ailton de Aquino Santos, Diogo Abry Guillen, Gilneu Francisco Astolfi Vivan, Izabela Moreira Correa, Nilton José Schneider David, Paulo Picchetti, Renato Dias de Brito Gomes e Rodrigo Alves Teixeira.

'''

    if valor_api != '':

        prompt = f"""
        Atue como um analista econômico experiente, especializado em analisar as atas do Copom, Comitê de Política Monetária
        do Banco Central do Brasil. A sua tarefa será analisar as atas que o Copom emite através de um arquivo de texto que
        será fornecido e produzir um resumo estruturado em Markdown para que os analistas possam entender melhor o que foi
        dito em cada reunião.
        
        O conteúdo do resumo deve conter:
        Um resumo principal do que foi discutido nas reuniões, contendo a data das reuniões, situação da taxa básica de juros,
        inflação e desemprego. Avalie também se a votação foi unânime ou não.
        Além disso, tópicos devem ser escritos a fim de especificar melhor cada assunto.
        1 - Situação do mercado interno.
        2 - Situação do mercado global.
        3 - Perspectivas gerais para as próximas reuniões
        Por fim, cite alguns destaques principais em forma de tópico.
        
        O texto a ser analisado é: {texto_copom}
        """
        response = model_ai.generate_content(prompt)
        pdf_response = markdown_2pdf(response.text)

        if pdf_response:
            st.download_button('Download resumo Copom .doc',pdf_response,file_name=f'Resumo Ata Copom {nro_ata_selecionada}.doc', mime='application/pdf')
        if pdf_response:
            st.download_button('Download resumo Copom .pdf',pdf_response,file_name=f'Resumo Ata Copom {nro_ata_selecionada}.pdf', mime='application/pdf')

if valor_api != '':
    st.write(response.text)

prompt_fed = f"""
Atue como um analista econômico experiente, especializado em analisar as atas do FOMC, do FED. 
A sua tarefa será analisar os Statments das reuniões do FOMC que emite através de um arquivo de texto que
será fornecido e produzir um resumo estruturado em Markdown para que os analistas possam entender melhor o que foi
dito em cada reunião.

O conteúdo do resumo deve conter:
Um resumo principal sobre o que foi discutido nas reuniões, contendo a data das reuniões, situação da taxa básica de juros dos EUA,
inflação e desemprego se houver. Faça um resumo bem feito com conectivos entre as frases ligando os assuntos tratados.
Além disso, tópicos devem ser escritos a fim de especificar melhor cada assunto.

1 - Situação do mercado interno americano.
2 - Situação do mercado global.
3 - Perspectivas gerais para as próximas reuniões

Por fim, cite alguns destaques principais em forma de tópico.

Dê mais foco do texto a ser redigido a partir de Decisions Regarding Monetary Policy Implementation

O resumo principal não deve conter mais do que 400 caracteres
Observação, nunca deixe de explicar o motivo de tal decisão. Cada ação tem um motivo, portanto explicite
O texto a ser analisado é: {texto_fed}
"""

prompt_fed_minutas = f"""
Atue como um analista econômico experiente, especializado em analisar as atas do FOMC, do FED. 
A sua tarefa será analisar os Statments das reuniões do FOMC que emite através de um arquivo de texto que
será fornecido e produzir um resumo estruturado em Markdown para que os analistas possam entender melhor o que foi
dito em cada reunião.

O conteúdo do resumo deve conter:
Um resumo principal sobre o que foi discutido nas reuniões, contendo a data das reuniões, situação da taxa básica de juros dos EUA,
inflação e desemprego se houver. Faça um resumo bem feito com conectivos entre as frases ligando os assuntos tratados.
Além disso, tópicos devem ser escritos a fim de especificar melhor cada assunto.

1 - Situação do mercado interno americano.
2 - Situação do mercado global.
3 - Perspectivas gerais para as próximas reuniões

Por fim, cite alguns destaques principais em forma de tópico.Estou p

Dê mais foco do texto a ser redigido a partir de Decisions Regarding Monetary Policy Implementation

O resumo principal inicial não deve conter mais do que 400 caracteres.
Além disso o restante do texto pode ser mais bem explicado não necessariamente só em tópicos

Observação, nunca deixe de explicar o motivo de tal decisão. Cada ação tem um motivo, portanto explicite
O texto a ser analisado é: {texto_fed}
Lembre-se, mantenha sempre a coesão do texto com um português adequado e pontuação correta, além de utilizar conectivos para cada frase.
"""


# Função para criar o conteúdo do arquivo .docx
def gerar_arquivo_docx(texto):
    # Cria o documento
    doc = docx.Document()
    doc.add_heading('Resumo Press Release FED', 0)

    # Adicionando conteúdo ao documento com caracteres especiais
    doc.add_paragraph(texto)

    # Salvar o documento em um objeto BytesIO (para evitar salvar em disco)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Reposiciona o ponteiro no início do buffer

    return buffer

response_fed = model_ai.generate_content(prompt_fed_minutas)
pdf_response_fed = markdown_2pdf(response_fed.text)
docx_fed = gerar_arquivo_docx(response_fed.text)

if pdf_response_fed:
    st.sidebar.download_button('Download resumo FED .doc', docx_fed, file_name=f'Resumo Press Release FED.docx',
                       mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
if pdf_response_fed:
    st.sidebar.download_button('Download resumo FED .pdf', pdf_response_fed, file_name=f'Resumo Press Release FED.pdf',
                       mime='application/pdf')

st.write(response_fed.text)
